import datetime

from aiogram import types
from aiogram.dispatcher import FSMContext
from aiogram.types import InputFile
from aiogram.utils import markdown

from calendar_google.calendar_api import GoogleCalendar
from db_api.dal.context_history_dal import ContextDAL
from db_api.dal.user_dal import UserDAL
from db_api.model import Context
from loader import dp, bot
from utils.openai_api import req_async_openai, req_nalog_api
from data.config import _key, _system
import docx

async def create_event(user_id, dict_mes):
    calendar_id = (await UserDAL.get(id=user_id)).calendar_id
    if calendar_id:
        if dict_mes[-1] != '}':
            dict_mes = dict_mes.split('}')[0]
        dict_mes = eval(dict_mes)

        where = dict_mes['где']
        _time = dict_mes['через сколько дней']
        _time_start = (datetime.datetime.now() + datetime.timedelta(days=float(_time)))
        _time_start_str = _time_start.strftime('%Y-%m-%d')
        _time_end = _time_start+ datetime.timedelta(days=1)
        _time_end_str = _time_end.strftime('%Y-%m-%d')

        event = {
            'summary': 'Встреча',
            'location': where,
            'description': 'Встреча',
            'start': {
                'date': _time_start_str,
            },
            'end': {
                'date': _time_end_str,
            },
        }
        obj = GoogleCalendar()
        event_add = obj.add_event(calendar_id, event)
        if event_add:
            await bot.send_message(user_id, 'Событие успешно добавлено')
        else:
            await bot.send_message(user_id, 'Произошла ошибка. Событие не добавлено')
    else:
        await bot.send_message(user_id, 'Вы еще не добавили свой гугл-календарь, чтобы его добавить отправьте команду /calendar')

async def create_document(user_id, response):
    # info
    customer = ''  # Введите название предприятия-заказчика
    customer_rs = ''  # Введите номер расчетного счета предприятия-заказчика
    customer_bik = ''  # Введите БИК банка предприятия-заказчика
    customer_fio = ''  # Введите ФИО представителя предприятия-заказчика для подписи
    customer_inn = ''  # Введите ИНН предприятия-заказчика

    executive = ''  # Введите название предприятия-исполнителя (Например, ИП Ивановиваниванович)
    executive_rs = ''  # Введите номер расчетного счета предприятия-исполнителя
    executive_bik = ''  # Введите БИК банка предприятия-исполнителя
    executive_fio = ''  # Введите ФИО представителя предприятия-исполнителя для подписи
    executive_inn = ''  # Введите ИНН предприятия-исполнителя

    if response[-1] != '}':
        response = response.split('}')[0]
    print(response)
    response = eval(response)
    print(response['с кем'])
    company = await req_nalog_api(response['с кем'])
    if not company['Count']:
        await bot.send_message(user_id, 'Компания не найдена. Попробуйте еще раз')
        return
    service = 'Продать рекламу'  # Кратко опишите услугу, которые обязан предоставить исполнитель
    date = datetime.datetime.now().strftime("%Y-%m-%d")
    price = response['цена']  # Введите сумму договора
    deadline = response['срок оплаты']  # Введите срок оплаты в днях (например, 10)

    replace_dict = {
        '{НазваниеКонтр}': company['items'][0]['ЮЛ']['НаимПолнЮЛ'],
        '{ИННКонтр}': company['items'][0]['ЮЛ']['ИНН'],
        '{РасчетныйСчетКонтр}': customer_rs,
        '{БИКБанкаКонтр}': customer_bik,
        '{ФИОКонтрДляПодписи}': customer_fio,

        '{ФИОИП}': executive,
        '{ИНН}': executive_inn,
        '{РасчетныйСчет}': executive_rs,
        '{БИК}': executive_bik,
        '{ФИОДляПодписи}': executive_fio,

        '{ДатаДокумента}': date,
        '{ОписаниеУслуг}': service,
        '{суммаДоговора}': str(price) + ' рублей',
        '{колвоДней}': deadline
    }



    doc = docx.Document("Договор_об_оказании_услуг.docx")

    paras = doc.paragraphs

    for i, k in replace_dict.items():
        for para in paras:
            para.text = para.text.replace(i, str(k))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace(i, str(k))

    doc.save('Заполненный_договор.docx')
    await bot.send_document(
        user_id, InputFile('Заполненный_договор.docx')
    )


@dp.message_handler(state=None, chat_type=[types.ChatType.PRIVATE])
async def bot_request(message: types.Message, state: FSMContext):
    _text = message.text
    user_id = message.chat.id

    # context = await ContextDAL.get(many=True, user_id=user_id)

    messages = list()
    messages.append({'role': 'system', 'content': _system()})
    # if context:
    #     for context_ in context[-10:]:
    #         messages.append({'role': context_.role, 'content': context_.content})
    messages.append({'role': "user", 'content': _text})
    # print(messages)

    response = await req_async_openai(messages=messages, _key=_key)

    response = response['choices'][0]['message']['content']
    # await ContextDAL.adds(
    #     Context(user_id=user_id, role='user', content=_text),
    #     Context(user_id=user_id, role='assistant', content=response)
    # )
    if '"запланировать встречу"' in response:
        await create_event(user_id, response)
    elif '"заполнить договор"' in response:
        await create_document(user_id, response)
    else:
        await bot.send_message(
            text=markdown.quote_html(response.replace('<br>', '\n'))[
                 :4096],
            chat_id=user_id,
            parse_mode=types.ParseMode.HTML
        )